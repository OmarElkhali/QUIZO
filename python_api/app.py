# -*- coding: utf-8 -*-
import os
import json
import logging
import requests
import re
from functools import wraps
from pypdf import PdfReader
import docx
import io
from flask import Flask, request, jsonify, g
from flask_cors import CORS
try:
    from flask_socketio import SocketIO, emit, join_room, leave_room
    SOCKETIO_AVAILABLE = True
except ImportError:
    SOCKETIO_AVAILABLE = False
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from dotenv import load_dotenv
from werkzeug.utils import secure_filename
import google.generativeai as genai
import firebase_admin
from firebase_admin import credentials, auth as firebase_auth

# Charger les variables d'environnement
load_dotenv()
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

app = Flask(__name__)
MAX_UPLOAD_SIZE_MB = int(os.getenv("MAX_UPLOAD_SIZE_MB", "10"))
MAX_AI_QUESTIONS_PER_QUIZ = int(os.getenv("MAX_AI_QUESTIONS_PER_QUIZ", "20"))
ALLOWED_UPLOAD_EXTENSIONS = {".pdf", ".docx", ".txt"}
app.config["MAX_CONTENT_LENGTH"] = MAX_UPLOAD_SIZE_MB * 1024 * 1024

# --- Firebase Admin SDK Initialization ---
_firebase_app = None
service_account_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
if service_account_path and os.path.isfile(service_account_path):
    try:
        cred = credentials.Certificate(service_account_path)
        _firebase_app = firebase_admin.initialize_app(cred)
        logging.getLogger(__name__).info("Firebase Admin SDK initialise avec succes")
    except Exception as e:
        logging.getLogger(__name__).warning(f"Firebase Admin SDK init echoue: {e}")
else:
    logging.getLogger(__name__).warning(
        "GOOGLE_APPLICATION_CREDENTIALS non defini ou fichier introuvable. "
        "La verification des tokens Firebase sera desactivee."
    )

# --- Rate Limiter ---
limiter = Limiter(
    app=app,
    key_func=get_remote_address,
    default_limits=["100 per minute"],
    storage_uri="memory://",
)

# --- CORS ---
configured_origins = os.getenv(
    "CORS_ORIGINS",
    "http://localhost:5173,http://127.0.0.1:5173,http://localhost:8080,http://127.0.0.1:8080",
).split(",")
local_dev_origins = [
    "http://localhost:5173",
    "http://127.0.0.1:5173",
    "http://localhost:8080",
    "http://127.0.0.1:8080",
]
ALLOWED_ORIGINS = list(dict.fromkeys(
    origin.strip()
    for origin in [*configured_origins, *local_dev_origins]
    if origin.strip()
))
CORS(app, resources={
    r"/*": {
        "origins": ALLOWED_ORIGINS,
        "methods": ["GET", "POST", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"],
        "max_age": 3600
    }
})

# Configuration du logging basee sur l'environnement
LOG_LEVEL = os.getenv("LOG_LEVEL", "INFO")
logging.basicConfig(
    level=getattr(logging, LOG_LEVEL),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)
logger = logging.getLogger(__name__)


# --- Firebase Auth Verification Decorator ---
def require_auth(f):
    """Decorator that verifies Firebase ID token from Authorization header.
    In dev mode (no Firebase Admin SDK), requests pass through with a warning.
    """
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # CORS preflight requests must pass through without auth
        if request.method == 'OPTIONS':
            return f(*args, **kwargs)

        if not _firebase_app:
            # Dev mode: no Firebase Admin SDK available, skip verification
            g.user_uid = None
            return f(*args, **kwargs)

        auth_header = request.headers.get('Authorization', '')
        if not auth_header.startswith('Bearer '):
            # In development, allow requests without auth (for testing)
            if os.getenv('FLASK_ENV') == 'development':
                logger.debug("Dev mode: requete sans token autorisee")
                g.user_uid = None
                return f(*args, **kwargs)
            return jsonify({'error': 'Token d\'authentification manquant'}), 401

        token = auth_header.replace('Bearer ', '')
        try:
            decoded_token = firebase_auth.verify_id_token(token)
            g.user_uid = decoded_token.get('uid')
            return f(*args, **kwargs)
        except firebase_auth.InvalidIdTokenError:
            return jsonify({'error': 'Token invalide'}), 401
        except firebase_admin.exceptions.FirebaseError as e:
            logger.error(f"Erreur Firebase Auth: {sanitize_error_message(str(e))}")
            return jsonify({'error': 'Erreur de verification du token'}), 401
        except Exception as e:
            logger.error(f"Erreur inattendue verification token: {sanitize_error_message(str(e))}")
            return jsonify({'error': 'Erreur de verification du token'}), 401

    return decorated_function


# Clés API depuis les variables d'environnement
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Configuration des modèles
OPENROUTER_MODEL = os.getenv("OPENROUTER_MODEL", "qwen/qwen3.6-flash")
GROQ_MODEL = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
GROQ_BASE_URL = os.getenv("GROQ_BASE_URL") or "https://api.groq.com/openai/v1"
DEFAULT_GROQ_MODEL = GROQ_MODEL

SUPPORTED_MODELS = {"gemini", "openrouter", "groq"}
PROVIDER_FALLBACK_ORDER = ["gemini", "openrouter", "groq"]


def sanitize_error_message(message):
    """Remove API key values from provider errors before logging or returning them."""
    if not message:
        return ""
    sanitized = re.sub(r"api_key:[A-Za-z0-9_\-]+", "api_key:[redacted]", str(message))
    sanitized = re.sub(r"AIza[0-9A-Za-z_\-]{20,}", "[redacted-google-api-key]", sanitized)
    sanitized = re.sub(r"sk-or-[0-9A-Za-z_\-]{20,}", "[redacted-openrouter-api-key]", sanitized)
    sanitized = re.sub(r"sk-[0-9A-Za-z_\-]{20,}", "[redacted-secret-key]", sanitized)
    sanitized = re.sub(r"gsk_[0-9A-Za-z_\-]{20,}", "[redacted-groq-api-key]", sanitized)
    return sanitized


# Configuration de l'API Gemini
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)
    logger.info("API Gemini configurée avec succès")
else:
    logger.warning("GEMINI_API_KEY n'est pas configurée - la génération avec Gemini sera désactivée")

# OpenRouter et Groq ont des clés de secours par défaut
logger.info("Modèles OpenRouter et Groq disponibles avec clés de secours configurées")



def get_upload_extension(filename):
    _, extension = os.path.splitext((filename or "").lower())
    return extension


def assert_allowed_upload(file):
    safe_name = secure_filename(file.filename or "")
    extension = get_upload_extension(safe_name)
    if not safe_name or extension not in ALLOWED_UPLOAD_EXTENSIONS:
        allowed = ", ".join(sorted(ALLOWED_UPLOAD_EXTENSIONS))
        raise ValueError(f"Format de fichier non supporte. Formats autorises: {allowed}")
    return safe_name, extension


def extract_text_from_file(file):
    """Extrait le texte des fichiers PDF/DOCX/TXT"""
    filename, extension = assert_allowed_upload(file)
    logger.debug(f"Début de l'extraction du texte pour le fichier: {filename}")
    
    try:
        if extension == '.pdf':
            logger.info(f"Extraction du texte du PDF: {filename}")
            pdf_reader = PdfReader(io.BytesIO(file.read()))
            logger.debug(f"PDF chargé avec {len(pdf_reader.pages)} pages")
            text = '\n'.join([page.extract_text() for page in pdf_reader.pages])
            logger.debug(f"Extraction PDF terminée: {len(text)} caractères extraits")
            return text if text.strip() else "Aucun texte détecté dans le PDF"
        
        elif extension == '.docx':
            logger.info(f"Extraction du texte du document Word: {filename}")
            doc = docx.Document(io.BytesIO(file.read()))
            logger.debug(f"Document Word chargé avec {len(doc.paragraphs)} paragraphes")
            text = '\n'.join([para.text for para in doc.paragraphs if para.text])
            logger.debug(f"Extraction Word terminée: {len(text)} caractères extraits")
            return text
        
        elif extension == '.txt':
            logger.info(f"Extraction du texte brut: {filename}")
            text = file.read().decode('utf-8', errors='ignore')
            logger.debug(f"Extraction TXT terminée: {len(text)} caractères extraits")
            return text
            
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte: {str(e)}")
        raise ValueError(f"Échec de l'extraction: {str(e)}")


@app.route('/api/extract-text', methods=['POST', 'OPTIONS'])
@limiter.limit("10 per minute")
@require_auth
def handle_text_extraction():
    """Endpoint pour l'extraction de texte depuis un fichier"""
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        logger.info("Requête d'extraction de texte reçue")
        
        if 'file' not in request.files:
            logger.warning("Aucun fichier fourni dans la requête")
            return jsonify({'error': 'Aucun fichier fourni'}), 400

        file = request.files['file']
        if not file or file.filename == '':
            logger.warning("Nom de fichier vide")
            return jsonify({'error': 'Nom de fichier vide'}), 400

        logger.info(f"Début de l'extraction pour le fichier: {file.filename}")
        text = extract_text_from_file(file)
        logger.info(f"Extraction réussie: {len(text)} caractères extraits")
        
        return jsonify({
            'text': text,
            'filename': file.filename,
            'characters': len(text)
        })

    except ValueError as e:
        logger.error(f"Erreur de format de fichier: {str(e)}")
        return jsonify({
            'error': f'Format de fichier non supporté ou fichier corrompu: {str(e)}',
            'details': str(e)
        }), 422
    except Exception as e:
        logger.error(f"Erreur inattendue lors de l'extraction: {str(e)}", exc_info=True)
        return jsonify({
            'error': 'Impossible d\'extraire le texte du fichier. Veuillez vérifier que le fichier est valide.',
            'details': str(e)
        }), 500


def validate_question(question):
    """Valide la structure d'une question générée"""
    required_fields = ['text', 'options', 'explanation', 'difficulty']
    if not all(field in question for field in required_fields):
        logger.warning(f"Question invalide: champs manquants - {question.get('text', 'Sans texte')}")
        return False
    if len(question['options']) != 4 or sum(opt.get('isCorrect', False) for opt in question['options']) != 1:
        logger.warning(f"Question invalide: problème avec les options - {question.get('text', 'Sans texte')}")
        return False
    return True


def normalize_manual_assistant_question(question, index, difficulty):
    """Normalise une proposition QCM pour le builder manuel."""
    if not isinstance(question, dict):
        return None

    text = str(question.get("text") or "").strip()
    raw_options = question.get("options")
    if not text or not isinstance(raw_options, list):
        return None

    options = []
    for option_index, option in enumerate(raw_options[:6]):
        if not isinstance(option, dict):
            continue
        option_text = str(option.get("text") or "").strip()
        if not option_text:
            continue
        options.append({
            "id": f"assistant_q{index + 1}_{chr(97 + len(options))}",
            "text": option_text,
            "isCorrect": bool(option.get("isCorrect")),
        })

    if len(options) < 2:
        return None

    if not any(option["isCorrect"] for option in options):
        options[0]["isCorrect"] = True

    if sum(1 for option in options if option["isCorrect"]) > 1:
        first_correct_seen = False
        for option in options:
            if option["isCorrect"] and not first_correct_seen:
                first_correct_seen = True
            else:
                option["isCorrect"] = False

    return {
        "id": f"assistant_q{index + 1}",
        "text": text,
        "options": options,
        "explanation": str(question.get("explanation") or "Explication a valider.").strip(),
        "difficulty": question.get("difficulty") if question.get("difficulty") in ["easy", "medium", "hard"] else difficulty,
        "points": int(question.get("points") or 1),
    }


def build_manual_assistant_prompt(data):
    action = str(data.get("action") or "generate_from_course").strip()
    course_text = str(data.get("courseText") or data.get("text") or "").strip()
    current_question = data.get("question") if isinstance(data.get("question"), dict) else {}
    num_questions = max(1, min(int(data.get("numQuestions") or 3), 10))
    difficulty = data.get("difficulty") if data.get("difficulty") in ["easy", "medium", "hard"] else "medium"
    language = str(data.get("language") or "fr").strip()[:8]

    return f"""
Tu es l'assistant pedagogique de QUIZO pour construire des QCM.
Utilise uniquement le contenu fourni par l'enseignant. Reponds en JSON valide, sans markdown.

ACTION: {action}
LANGUE: {language}
DIFFICULTE: {difficulty}
NOMBRE DE PROPOSITIONS: {num_questions}

COURS:
{course_text[:7000]}

QUESTION ACTUELLE EVENTUELLE:
{json.dumps(current_question, ensure_ascii=False)[:2500]}

FORMAT JSON STRICT:
{{
  "summary": "resume court de ce que tu as fait",
  "issues": ["probleme detecte ou conseil"],
  "suggestions": [
    {{
      "text": "Question QCM claire",
      "options": [
        {{"text": "Bonne reponse", "isCorrect": true}},
        {{"text": "Distracteur plausible", "isCorrect": false}},
        {{"text": "Distracteur plausible", "isCorrect": false}},
        {{"text": "Distracteur plausible", "isCorrect": false}}
      ],
      "explanation": "Explication concise basee sur le cours",
      "difficulty": "{difficulty}",
      "points": 1
    }}
  ]
}}

CONTRAINTES:
- Une seule option correcte par question.
- Les distracteurs doivent etre plausibles mais faux.
- Evite les formulations ambigues.
- Si le cours est insuffisant, renvoie au moins un probleme dans "issues".
"""


@app.route('/api/manual-assistant', methods=['POST', 'OPTIONS'])
@limiter.limit("10 per day;5 per minute")
@require_auth
def manual_assistant():
    """Assistant OpenRouter pour le builder de quiz manuel."""
    if request.method == 'OPTIONS':
        return '', 204

    try:
        data = request.get_json() or {}
        course_text = str(data.get("courseText") or data.get("text") or "").strip()
        action = str(data.get("action") or "generate_from_course").strip()
        difficulty = data.get("difficulty") if data.get("difficulty") in ["easy", "medium", "hard"] else "medium"
        num_questions = max(1, min(int(data.get("numQuestions") or 3), 10))

        if action == "generate_from_course" and len(course_text) < 80:
            return jsonify({
                "error": "Ajoutez un cours plus complet avant de generer des questions.",
                "details": "Le contenu doit contenir au moins 80 caracteres.",
            }), 400

        if not OPENROUTER_API_KEY:
            fallback = generate_fallback_questions(num_questions, difficulty, course_text)
            return jsonify({
                "summary": "OpenRouter n'est pas configure. Propositions locales de secours generees.",
                "issues": ["Configurez OPENROUTER_API_KEY pour activer l'assistant complet."],
                "suggestions": [
                    normalize_manual_assistant_question(question, index, difficulty)
                    for index, question in enumerate(fallback)
                ],
                "provider": "local-fallback",
                "model": "local-fallback",
                "fallback": True,
            }), 200

        prompt = build_manual_assistant_prompt({
            **data,
            "courseText": course_text,
            "difficulty": difficulty,
            "numQuestions": num_questions,
        })
        content = generate_with_openrouter(prompt)

        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if not json_match:
            raise ValueError("Aucun JSON trouve dans la reponse OpenRouter")

        payload = json.loads(json_match.group())
        raw_suggestions = payload.get("suggestions", [])
        if not isinstance(raw_suggestions, list):
            raw_suggestions = []

        suggestions = [
            normalized
            for index, question in enumerate(raw_suggestions[:num_questions])
            for normalized in [normalize_manual_assistant_question(question, index, difficulty)]
            if normalized is not None
        ]

        if not suggestions:
            raise ValueError("OpenRouter n'a pas retourne de proposition QCM valide")

        issues = payload.get("issues", [])
        if not isinstance(issues, list):
            issues = []

        return jsonify({
            "summary": str(payload.get("summary") or "Propositions generees avec OpenRouter."),
            "issues": [str(issue) for issue in issues[:6]],
            "suggestions": suggestions,
            "provider": "openrouter",
            "model": OPENROUTER_MODEL,
            "fallback": False,
        })
    except ValueError as e:
        safe_error = sanitize_error_message(str(e))
        logger.warning(f"Assistant manuel indisponible: {safe_error}")
        if 'course_text' in locals() and str(course_text).strip():
            fallback = generate_fallback_questions(num_questions if 'num_questions' in locals() else 3, difficulty if 'difficulty' in locals() else "medium", course_text)
            return jsonify({
                "summary": "OpenRouter est indisponible. Propositions locales de secours generees.",
                "issues": [safe_error],
                "suggestions": [
                    normalize_manual_assistant_question(question, index, difficulty if 'difficulty' in locals() else "medium")
                    for index, question in enumerate(fallback)
                ],
                "provider": "local-fallback",
                "model": "local-fallback",
                "fallback": True,
            }), 200
        return jsonify({
            "error": "Assistant OpenRouter indisponible. Veuillez reessayer ou verifier la configuration.",
            "details": safe_error,
        }), 503
    except Exception as e:
        safe_error = sanitize_error_message(str(e))
        logger.error(f"Erreur inattendue assistant manuel: {safe_error}", exc_info=True)
        return jsonify({
            "error": "Erreur inattendue pendant l'assistance de creation.",
            "details": safe_error,
        }), 500


@app.route('/api/generate', methods=['POST', 'OPTIONS'])
@limiter.limit("20 per minute")
@require_auth
def generate_quiz():
    """Endpoint principal pour la génération de quiz"""
    if request.method == 'OPTIONS':
        return '', 204
    
    try:
        logger.info("Requête de génération de quiz reçue")
        
        # Gestion des entrées multiples
        data = {}
        if 'file' in request.files:
            file = request.files['file']
            logger.info(f"Extraction du texte à partir du fichier: {file.filename}")
            text = extract_text_from_file(file)
            data = {
                'numQuestions': request.form.get('numQuestions'),
                'difficulty': request.form.get('difficulty'),
                'modelType': request.form.get('modelType'),
            }
        else:
            data = request.get_json() or {}
            text = data.get('text', '')
            logger.info(f"Texte reçu directement: {len(text)} caractères")

        # Validation des paramètres
        if not text or not text.strip():
            logger.warning("Aucun contenu textuel fourni pour la génération")
            return jsonify({'error': 'Aucun contenu fourni'}), 400

        try:
            num_questions = int(data.get('numQuestions', 5))
            if num_questions < 1 or num_questions > MAX_AI_QUESTIONS_PER_QUIZ:
                return jsonify({
                    'error': f'Le nombre de questions doit etre entre 1 et {MAX_AI_QUESTIONS_PER_QUIZ}'
                }), 400
        except (ValueError, TypeError):
            return jsonify({'error': 'Nombre de questions invalide'}), 400

        difficulty = data.get('difficulty', 'medium')
        if difficulty not in ['easy', 'medium', 'hard']:
            difficulty = 'medium'

        model_type = (data.get('modelType') or 'gemini').strip().lower()
        if model_type not in SUPPORTED_MODELS:
            return jsonify({
                'error': f"Modele IA non supporte: {model_type}",
                'supportedModels': sorted(SUPPORTED_MODELS),
            }), 400

        api_key = data.get('apiKey', '')
        
        logger.info(f"Paramètres de génération: {num_questions} questions, difficulté {difficulty}, modèle {model_type}")

        # Vérifier que la clé API nécessaire est disponible
        if model_type == 'gemini' and not GEMINI_API_KEY:
            return jsonify({'error': 'Clé API Gemini non configurée dans python_api/.env'}), 503

        # Construction du prompt
        logger.debug(f"Construction du prompt avec {len(text[:5000])} caractères de texte")
        prompt = f"""
        Généree {num_questions} questions QCM complexes en français selon ces règles STRICTES :

        TEXTE SOURCE :
        {text[:5000]}

        FORMAT JSON REQUIS :
        {{
          "questions": [
            {{
              "text": "Question...",
              "options": [
                {{"text": "...", "isCorrect": true}},
                {{"text": "...", "isCorrect": false}},
                {{"text": "...", "isCorrect": false}},
                {{"text": "...", "isCorrect": false}}
              ],
              "explanation": "Explication basée sur le texte",
              "difficulty": "{difficulty}"
            }}
          ]
        }}

        CONTRAINTES :
        - Une seule réponse correcte par question
        - Les options doivent être plausibles
        - Les explications doivent citer le texte
        - Ne pas inclure de markdown
        {f"- Contraintes supplementaires: {data.get('additionalInfo')}" if data.get('additionalInfo') else ""}
        """
        logger.debug("Prompt construit avec succès")

        content = None
        used_provider = model_type
        provider_errors = []
        for provider in get_provider_attempt_order(model_type):
            try:
                logger.info(f"Tentative de generation avec {provider}")
                content = generate_with_provider(provider, prompt, api_key)
                used_provider = provider
                break
            except ValueError as e:
                safe_error = sanitize_error_message(str(e))
                provider_errors.append(f"{provider}: {safe_error}")
                logger.warning(f"Provider {provider} indisponible: {safe_error}")

        if not content:
            raise ValueError("Service IA indisponible: " + " | ".join(provider_errors))

        logger.debug(f"Contenu brut reçu: {len(content)} caractères")
        logger.debug(f"Aperçu du contenu: {content[:200]}")
        
        # Chercher le JSON dans la réponse
        json_match = re.search(r'\{.*\}', content, re.DOTALL)
        if not json_match:
            logger.error("Aucun JSON trouvé dans la réponse")
            logger.error(f"Contenu complet: {content}")
            raise ValueError("Aucun JSON trouvé dans la réponse de l'IA")
        
        try:
            quiz_data = json.loads(json_match.group())
        except json.JSONDecodeError as e:
            logger.error(f"Erreur de décodage JSON: {e}")
            logger.error(f"Contenu extrait: {json_match.group()}")
            raise ValueError(f"JSON mal formé retourné par l'IA: {e}")

        # Validation et formatage des questions
        questions = quiz_data.get('questions', [])
        if not isinstance(questions, list) or len(questions) == 0:
            logger.warning("Aucune question trouvée dans le JSON généré")
            raise ValueError("Le JSON de l'IA ne contient pas de liste de questions")

        valid_questions = []
        for index, q in enumerate(questions):
            # Normaliser et vérifier la validité minimale
            if not isinstance(q, dict) or 'text' not in q or 'options' not in q:
                continue

            q_text = str(q.get('text', '')).strip()
            raw_options = q.get('options', [])

            if not q_text or not isinstance(raw_options, list) or len(raw_options) < 2:
                continue

            # Valider les options et injecter les IDs
            options = []
            for option_index, opt in enumerate(raw_options[:4]):
                if not isinstance(opt, dict):
                    continue
                opt_text = str(opt.get('text', '')).strip()
                if not opt_text:
                    continue
                options.append({
                    'id': f"q{len(valid_questions) + 1}_{chr(97 + option_index)}",
                    'text': opt_text,
                    'isCorrect': bool(opt.get('isCorrect', False))
                })

            if len(options) < 2:
                continue

            # S'assurer qu'au moins une réponse est correcte
            if not any(opt['isCorrect'] for opt in options):
                options[0]['isCorrect'] = True

            # S'assurer d'une seule réponse correcte
            if sum(1 for opt in options if opt['isCorrect']) > 1:
                first_correct = False
                for opt in options:
                    if opt['isCorrect'] and not first_correct:
                        first_correct = True
                    else:
                        opt['isCorrect'] = False

            valid_questions.append({
                'id': f"gen_q{len(valid_questions) + 1}",
                'text': q_text,
                'options': options,
                'explanation': str(q.get('explanation') or 'Explication à retenir.').strip(),
                'difficulty': q.get('difficulty') if q.get('difficulty') in ['easy', 'medium', 'hard'] else difficulty
            })

        logger.info(f"Génération terminée: {len(valid_questions)} questions valides sur {num_questions} demandées")

        if len(valid_questions) == 0:
            raise ValueError("L'IA n'a produit aucune question syntaxiquement correcte.")

        # Compléter si insuffisant
        warning = None
        if len(valid_questions) < num_questions:
            warning = f"Seulement {len(valid_questions)} questions valides ont été générées sur {num_questions} demandées."
            logger.warning(warning)
            # Génération locale fallback
            fallback = generate_fallback_questions(num_questions - len(valid_questions), difficulty, text)
            valid_questions.extend(fallback)

        return jsonify({
            'questions': valid_questions[:num_questions],
            'warning': warning,
            'provider': used_provider,
            'fallback_used': len(valid_questions) > num_questions
        })

    except ValueError as e:
        safe_error = sanitize_error_message(str(e))
        logger.error(f"Erreur fonctionnelle lors de la génération: {safe_error}")
        return jsonify({
            'error': 'Impossible de générer le quiz à partir de ce texte.',
            'details': safe_error
        }), 422
    except Exception as e:
        safe_error = sanitize_error_message(str(e))
        logger.error(f"Erreur inattendue lors de la génération: {safe_error}", exc_info=True)
        return jsonify({
            "error": "Une erreur inattendue s'est produite lors de la génération des questions. Veuillez réessayer.",
            "details": safe_error
        }), 500


def generate_with_gemini(prompt):
    try:
        logger.info("Utilisation de l'API Gemini avec le SDK officiel")
        model = genai.GenerativeModel('gemini-2.5-flash')
        response = model.generate_content(prompt)
        content = response.text
        logger.info(f"Contenu extrait de Gemini: {len(content)} caracteres")
        return content
    except Exception as e:
        safe_error = sanitize_error_message(str(e))
        logger.error(f"Erreur lors de l'appel a l'API Gemini: {safe_error}")
        raise ValueError(f"Service Gemini indisponible: {safe_error}")


def generate_with_chat_completions_api(provider_name, base_url, api_key, model, prompt, extra_headers=None):
    """Genere du contenu avec une API de chat completions compatible."""
    if not api_key:
        raise ValueError(f"Cle API {provider_name} non configuree")

    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        **(extra_headers or {}),
    }
    endpoint = f"{base_url.rstrip('/')}/chat/completions"

    try:
        logger.info(f"Envoi de la requete a {provider_name} avec le modele {model}")
        response = requests.post(
            endpoint,
            headers=headers,
            json={
                "model": model,
                "messages": [
                    {
                        "role": "system",
                        "content": "Tu es un expert en creation de quiz educatifs. Reponds uniquement avec un JSON valide.",
                    },
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.2,
            },
            timeout=90,
        )
        response.raise_for_status()
        payload = response.json()
        return payload["choices"][0]["message"]["content"]
    except requests.exceptions.RequestException as e:
        safe_error = sanitize_error_message(str(e))
        if getattr(e, "response", None) is not None:
            safe_error = f"{safe_error} | {sanitize_error_message(e.response.text[:800])}"
        logger.error(f"Erreur lors de l'appel a {provider_name}: {safe_error}")
        raise ValueError(f"Service {provider_name} indisponible: {safe_error}")
    except (KeyError, IndexError, TypeError, ValueError) as e:
        safe_error = sanitize_error_message(str(e))
        logger.error(f"Reponse invalide de {provider_name}: {safe_error}")
        raise ValueError(f"Reponse invalide de {provider_name}: {safe_error}")


def generate_with_openrouter(prompt, user_api_key=None):
    BACKUP_OPENROUTER_KEYS = [
        k.strip() for k in os.getenv("BACKUP_OPENROUTER_KEYS", "").split(",") if k.strip()
    ]
    keys_to_try = []
    if user_api_key:
        keys_to_try.append(user_api_key)
    if OPENROUTER_API_KEY and OPENROUTER_API_KEY != "your_openrouter_api_key_here":
        keys_to_try.append(OPENROUTER_API_KEY)
    for key in BACKUP_OPENROUTER_KEYS:
        if key not in keys_to_try:
            keys_to_try.append(key)
            
    errors = []
    for i, key in enumerate(keys_to_try):
        try:
            logger.info(f"OpenRouter: Tentative avec la clé index {i}")
            return generate_with_chat_completions_api(
                "OpenRouter",
                "https://openrouter.ai/api/v1",
                key,
                OPENROUTER_MODEL,
                prompt,
                {
                    "HTTP-Referer": os.getenv("APP_PUBLIC_URL", "http://localhost:5173"),
                    "X-Title": "QUIZO",
                },
            )
        except Exception as e:
            logger.warning(f"Clé OpenRouter index {i} a échoué: {sanitize_error_message(str(e))}")
            errors.append(f"Key {i}: {str(e)}")
            
    raise ValueError("Toutes les clés OpenRouter ont échoué. Détails: " + " | ".join(errors))


def generate_with_groq(prompt, user_api_key=None):
    BACKUP_GROQ_KEYS = [
        k.strip() for k in os.getenv("BACKUP_GROQ_KEYS", "").split(",") if k.strip()
    ]
    keys_to_try = []
    if user_api_key:
        keys_to_try.append(user_api_key)
    if GROQ_API_KEY and GROQ_API_KEY != "your_groq_api_key_here":
        keys_to_try.append(GROQ_API_KEY)
    for key in BACKUP_GROQ_KEYS:
        if key not in keys_to_try:
            keys_to_try.append(key)
            
    errors = []
    for i, key in enumerate(keys_to_try):
        try:
            logger.info(f"Groq: Tentative avec la clé index {i}")
            return generate_with_chat_completions_api(
                "Groq",
                GROQ_BASE_URL,
                key,
                GROQ_MODEL,
                prompt,
            )
        except Exception as e:
            logger.warning(f"Clé Groq index {i} a échoué: {sanitize_error_message(str(e))}")
            errors.append(f"Key {i}: {str(e)}")
            
    raise ValueError("Toutes les clés Groq ont échoué. Détails: " + " | ".join(errors))





def generate_fallback_questions(num, difficulty, source_text=""):
    """Genere des questions de secours a partir du texte extrait."""
    logger.info(f"Generation fallback de {num} questions avec difficulte {difficulty}")

    cleaned_text = re.sub(r"\s+", " ", source_text or "").strip()
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?؟])\s+", cleaned_text)
        if len(sentence.strip()) > 35
    ]

    if not sentences and cleaned_text:
        sentences = [
            cleaned_text[index:index + 220].strip()
            for index in range(0, min(len(cleaned_text), num * 220), 220)
            if cleaned_text[index:index + 220].strip()
        ]

    generic_fragments = [
        "Le document presente une idee principale liee au sujet etudie.",
        "Le passage insiste sur les notions essentielles du cours.",
        "Le contenu explique un concept important a retenir.",
        "Le texte donne des elements utiles pour comprendre le chapitre.",
    ]

    if not sentences:
        sentences = generic_fragments

    questions = []
    for i in range(num):
        qid = f"q{i + 1}"
        correct = sentences[i % len(sentences)][:180]
        topic = correct[:80]
        distractors = [
            generic_fragments[(i + 1) % len(generic_fragments)],
            generic_fragments[(i + 2) % len(generic_fragments)],
            generic_fragments[(i + 3) % len(generic_fragments)],
        ]

        questions.append({
            "id": qid,
            "text": f"D'apres le document, quel extrait correspond le mieux a cette idee : {topic} ?",
            "options": [
                {"id": f"{qid}_a", "text": correct, "isCorrect": True},
                {"id": f"{qid}_b", "text": distractors[0], "isCorrect": False},
                {"id": f"{qid}_c", "text": distractors[1], "isCorrect": False},
                {"id": f"{qid}_d", "text": distractors[2], "isCorrect": False},
            ],
            "explanation": "Question generee localement a partir du texte extrait, car le service IA externe est indisponible.",
            "difficulty": difficulty,
        })

    return questions


def get_provider_configured_state():
    return {
        "gemini": bool(GEMINI_API_KEY),
        "openrouter": True,
        "groq": True,
    }


def get_provider_model(provider):
    if provider == "gemini":
        return "gemini-2.5-flash"
    if provider == "openrouter":
        return OPENROUTER_MODEL
    if provider == "groq":
        return GROQ_MODEL
    return "unknown"


def get_provider_attempt_order(requested_model):
    """Return provider attempts in the free-beta fallback order."""
    if requested_model in PROVIDER_FALLBACK_ORDER:
        return [requested_model] + [
            provider for provider in PROVIDER_FALLBACK_ORDER if provider != requested_model
        ]
    return list(PROVIDER_FALLBACK_ORDER)


def assert_provider_ready(provider_name):
    states = get_provider_configured_state()
    if provider_name not in states:
        raise ValueError(f"Fournisseur IA inconnu: {provider_name}")
    if not states[provider_name]:
        raise ValueError(f"Le fournisseur {provider_name} n'est pas configuré (clé manquante)")


def generate_with_provider(model_type, prompt, api_key=None):
    """Call one configured provider without exposing provider credentials."""
    assert_provider_ready(model_type)
    if model_type == "gemini":
        return generate_with_gemini(prompt)
    if model_type == "openrouter":
        return generate_with_openrouter(prompt, api_key)
    if model_type == "groq":
        return generate_with_groq(prompt, api_key)
    raise ValueError(f"Modele IA non supporte: {model_type}")


@app.route('/api/health', methods=['GET'])
def health_check():
    """Endpoint de verification de l'etat du service"""
    services = get_provider_configured_state()

    return jsonify({
        "status": "ok",
        "version": "1.1.0",
        "services": services,
        "models": {
            provider: get_provider_model(provider)
            for provider in sorted(SUPPORTED_MODELS)
        },
        "groq": True
    })


@app.route('/api/status', methods=['GET'])
def status_check():
    """Status detaille — utile pour le monitoring"""
    services = get_provider_configured_state()
    providers_list = [p for p, configured in services.items() if configured]
    return jsonify({
        "status": "ok",
        "version": "1.1.0",
        "providers_configured": sorted(providers_list),
        "providers_count": len(providers_list),
        "firebase_admin_configured": _firebase_app is not None,
        "firebase_project_id": os.getenv("FIREBASE_PROJECT_ID", "not-set"),
        "environment": os.getenv("FLASK_ENV", "development"),
        "rate_limiting": True,
        "auth_verification": _firebase_app is not None,
    })


@app.route('/api/providers', methods=['GET'])
def providers_check():
    """Expose les fournisseurs IA connus sans divulguer les secrets."""
    services = get_provider_configured_state()
    return jsonify({
        "providers": {
            provider: {
                "configured": services.get(provider, False),
                "model": get_provider_model(provider),
            }
            for provider in sorted(SUPPORTED_MODELS)
        }
    })


# --- Request logging middleware ---
import time as _time

@app.before_request
def _log_request_start():
    request._start_time = _time.time()

@app.after_request
def _log_request_end(response):
    duration = round((_time.time() - getattr(request, '_start_time', _time.time())) * 1000)
    if request.path != '/api/health':  # don't spam health check logs
        logger.info(f"{request.method} {request.path} -> {response.status_code} ({duration}ms)")
    return response


# --- Live Session (Multiplayer Quiz) ---
try:
    from live_session import LiveSessionManager
    session_manager = LiveSessionManager()
    logger.info("LiveSessionManager charge avec succes")
except ImportError:
    session_manager = None
    logger.warning("live_session module non disponible")

if SOCKETIO_AVAILABLE:
    socketio = SocketIO(
        app,
        cors_allowed_origins=ALLOWED_ORIGINS,
        async_mode='threading',
        logger=False,
        engineio_logger=False,
    )
else:
    socketio = None


@app.route('/api/live/create', methods=['POST', 'OPTIONS'])
@require_auth
def create_live_session():
    """Creer une session de quiz en temps reel."""
    if request.method == 'OPTIONS':
        return '', 204
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Donnees JSON requises'}), 400

    title = data.get('title', 'Quiz sans titre')
    questions = data.get('questions', [])
    creator_id = getattr(g, 'user_uid', None) or 'anonymous'

    if not questions:
        return jsonify({'error': 'Au moins une question requise'}), 400

    try:
        session = session_manager.create_session(title, questions, creator_id)
        return jsonify({
            'sessionId': session.session_id,
            'accessCode': session.access_code,
            'quizTitle': session.quiz_title,
            'totalQuestions': session.total_questions,
            'status': session.status.value,
        })
    except Exception as e:
        logger.error(f"Erreur creation session live: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/live/join', methods=['POST', 'OPTIONS'])
def join_live_session():
    """Rejoindre une session avec un code d'acces."""
    if request.method == 'OPTIONS':
        return '', 204
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Donnees JSON requises'}), 400

    access_code = data.get('accessCode', '').strip().upper()
    player_name = data.get('playerName', 'Joueur').strip()

    if not access_code:
        return jsonify({'error': 'Code d\'acces requis'}), 400

    try:
        session, player = session_manager.join_session(access_code, player_name)

        if socketio:
            socketio.emit('player_joined', player.to_dict(), room=session.session_id)

        return jsonify({
            'sessionId': session.session_id,
            'player': player.to_dict(),
            'session': session.to_dict(),
        })
    except ValueError as e:
        return jsonify({'error': str(e)}), 400
    except Exception as e:
        logger.error(f"Erreur jonction session live: {e}")
        return jsonify({'error': str(e)}), 500


@app.route('/api/live/session/<session_id>', methods=['GET'])
def get_live_session(session_id):
    """Obtenir l'etat d'une session."""
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    session = session_manager.get_session(session_id)
    if not session:
        return jsonify({'error': 'Session non trouvee'}), 404

    return jsonify({
        'session': session.to_dict(),
        'players': [p.to_dict() for p in session.players.values()],
        'leaderboard': session_manager.get_leaderboard(session_id),
    })


@app.route('/api/live/session/<session_id>/start', methods=['POST', 'OPTIONS'])
@require_auth
def start_live_session(session_id):
    """Demarrer une session."""
    if request.method == 'OPTIONS':
        return '', 204
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    try:
        session = session_manager.start_session(session_id)

        if socketio:
            socketio.emit('session_started', session.to_dict(), room=session_id)

        return jsonify({'session': session.to_dict()})
    except ValueError as e:
        return jsonify({'error': str(e)}), 400


@app.route('/api/live/session/<session_id>/next', methods=['POST', 'OPTIONS'])
@require_auth
def next_live_question(session_id):
    """Passer a la question suivante."""
    if request.method == 'OPTIONS':
        return '', 204
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    question = session_manager.advance_to_question(session_id)
    session = session_manager.get_session(session_id)

    if not question:
        # Quiz completed
        if socketio and session:
            socketio.emit('session_completed', {
                'session': session.to_dict(),
                'leaderboard': session_manager.get_leaderboard(session_id),
            }, room=session_id)
        return jsonify({
            'completed': True,
            'leaderboard': session_manager.get_leaderboard(session_id),
        })

    question_data = {
        'id': question.id,
        'text': question.text,
        'options': question.options,
        'timeLimit': question.time_limit,
        'points': question.points,
        'questionIndex': session.current_question_index if session else 0,
    }

    if socketio:
        socketio.emit('new_question', question_data, room=session_id)

    return jsonify({'question': question_data, 'completed': False})


@app.route('/api/live/session/<session_id>/answer', methods=['POST', 'OPTIONS'])
def submit_live_answer(session_id):
    """Soumettre une reponse."""
    if request.method == 'OPTIONS':
        return '', 204
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Donnees JSON requises'}), 400

    try:
        answer = session_manager.submit_answer(
            session_id,
            data.get('playerId', ''),
            data.get('questionId', ''),
            data.get('selectedOptionId', ''),
        )

        answer_data = {
            'playerId': answer.player_id,
            'questionId': answer.question_id,
            'isCorrect': answer.is_correct,
            'pointsEarned': answer.points_earned,
            'timeSpent': answer.time_spent,
        }

        if socketio:
            socketio.emit('answer_submitted', answer_data, room=session_id)

        return jsonify(answer_data)
    except ValueError as e:
        return jsonify({'error': str(e)}), 400


@app.route('/api/live/session/<session_id>/leaderboard', methods=['GET'])
def get_live_leaderboard(session_id):
    """Obtenir le classement."""
    if not session_manager:
        return jsonify({'error': 'Live sessions non disponibles'}), 503

    return jsonify({
        'leaderboard': session_manager.get_leaderboard(session_id),
    })


# --- SocketIO Event Handlers ---
if SOCKETIO_AVAILABLE and socketio:
    @socketio.on('join_session')
    def handle_join_session(data):
        session_id = data.get('sessionId', '')
        join_room(session_id)
        emit('joined_room', {'sessionId': session_id})

    @socketio.on('leave_session')
    def handle_leave_session(data):
        session_id = data.get('sessionId', '')
        leave_room(session_id)


if __name__ == '__main__':
    logger.info("Demarrage du serveur Flask sur le port 5000")
    if SOCKETIO_AVAILABLE and socketio:
        logger.info("Flask-SocketIO actif - WebSocket disponible")
        socketio.run(app, host='0.0.0.0', port=5000, debug=False, allow_unsafe_werkzeug=True)
    else:
        logger.info("Mode REST uniquement (Flask-SocketIO non installe)")
        app.run(host='0.0.0.0', port=5000, debug=False)
